"""
Document loader utilities for Easy RAG System.
This module provides functionality for loading documents into vector databases.
"""

import os
import importlib
import subprocess
import sys
from typing import List, Dict, Any, Optional, Tuple

class DocumentLoader:
    """Base class for document loaders"""
    
    name = "Base Loader"
    description = "Base document loader class"
    supported_extensions = []
    required_packages = []
    
    @classmethod
    def get_name(cls) -> str:
        """Get the name of the loader"""
        return cls.name
    
    @classmethod
    def get_description(cls) -> str:
        """Get the description of the loader"""
        return cls.description
    
    @classmethod
    def get_supported_extensions(cls) -> List[str]:
        """Get the supported file extensions"""
        return cls.supported_extensions
    
    @classmethod
    def get_required_packages(cls) -> List[str]:
        """Get the required packages for this loader"""
        return cls.required_packages
    
    @classmethod
    def check_dependencies(cls) -> Tuple[bool, List[str]]:
        """
        Check if all required dependencies are installed
        Returns: (all_installed, missing_packages)
        """
        missing = []
        for package in cls.required_packages:
            try:
                importlib.import_module(package.split('==')[0])
            except ImportError:
                missing.append(package)
        
        return len(missing) == 0, missing
    
    @classmethod
    def install_dependencies(cls) -> Tuple[bool, str]:
        """
        Install required dependencies
        Returns: (success, message)
        """
        _, missing = cls.check_dependencies()
        if not missing:
            return True, "All dependencies already installed"
        
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install"] + missing)
            return True, f"Successfully installed: {', '.join(missing)}"
        except subprocess.CalledProcessError as e:
            return False, f"Failed to install dependencies: {str(e)}"
    
    def load_document(self, file_path: str) -> Dict[str, Any]:
        """
        Load a document from a file path
        Returns: Document content and metadata
        """
        raise NotImplementedError("Subclasses must implement this method")


class TextLoader(DocumentLoader):
    """Loader for plain text files"""
    
    name = "Text Loader"
    description = "Loads plain text files (.txt)"
    supported_extensions = ['.txt']
    required_packages = []
    
    def load_document(self, file_path: str) -> Dict[str, Any]:
        """Load a text document"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return {
            'content': content,
            'metadata': {
                'source': file_path,
                'file_type': 'text',
                'lines': content.count('\n') + 1,
                'characters': len(content)
            }
        }


class PDFLoader(DocumentLoader):
    """Loader for PDF files"""
    
    name = "PDF Loader"
    description = "Loads PDF documents (.pdf)"
    supported_extensions = ['.pdf']
    required_packages = ['PyPDF2']
    
    def load_document(self, file_path: str) -> Dict[str, Any]:
        """Load a PDF document"""
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            
            content = ""
            for page in reader.pages:
                content += page.extract_text() + "\n\n"
            
            return {
                'content': content,
                'metadata': {
                    'source': file_path,
                    'file_type': 'pdf',
                    'pages': len(reader.pages),
                    'characters': len(content)
                }
            }
        except ImportError:
            raise ImportError("PyPDF2 is required for PDF loading. Please install it first.")


class DocxLoader(DocumentLoader):
    """Loader for DOCX files"""
    
    name = "DOCX Loader"
    description = "Loads Microsoft Word documents (.docx)"
    supported_extensions = ['.docx']
    required_packages = ['python-docx']
    
    def load_document(self, file_path: str) -> Dict[str, Any]:
        """Load a DOCX document"""
        try:
            import docx
            doc = docx.Document(file_path)
            
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            return {
                'content': content,
                'metadata': {
                    'source': file_path,
                    'file_type': 'docx',
                    'paragraphs': len(doc.paragraphs),
                    'characters': len(content)
                }
            }
        except ImportError:
            raise ImportError("python-docx is required for DOCX loading. Please install it first.")


class CSVLoader(DocumentLoader):
    """Loader for CSV files"""
    
    name = "CSV Loader"
    description = "Loads CSV files (.csv)"
    supported_extensions = ['.csv']
    required_packages = ['pandas']
    
    def load_document(self, file_path: str) -> Dict[str, Any]:
        """Load a CSV document"""
        try:
            import pandas as pd
            df = pd.read_csv(file_path)
            
            # Convert DataFrame to string representation
            content = df.to_string()
            
            return {
                'content': content,
                'metadata': {
                    'source': file_path,
                    'file_type': 'csv',
                    'rows': len(df),
                    'columns': len(df.columns),
                    'headers': df.columns.tolist()
                }
            }
        except ImportError:
            raise ImportError("pandas is required for CSV loading. Please install it first.")


class JSONLoader(DocumentLoader):
    """Loader for JSON files"""
    
    name = "JSON Loader"
    description = "Loads JSON files (.json)"
    supported_extensions = ['.json']
    required_packages = []
    
    def load_document(self, file_path: str) -> Dict[str, Any]:
        """Load a JSON document"""
        import json
        
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Convert JSON to string representation
        content = json.dumps(data, indent=2)
        
        metadata = {
            'source': file_path,
            'file_type': 'json',
            'characters': len(content)
        }
        
        # Add structure info
        if isinstance(data, dict):
            metadata['structure'] = 'object'
            metadata['keys'] = list(data.keys())
        elif isinstance(data, list):
            metadata['structure'] = 'array'
            metadata['items'] = len(data)
        
        return {
            'content': content,
            'metadata': metadata
        }


class MarkdownLoader(DocumentLoader):
    """Loader for Markdown files"""
    
    name = "Markdown Loader"
    description = "Loads Markdown files (.md)"
    supported_extensions = ['.md']
    required_packages = []
    
    def load_document(self, file_path: str) -> Dict[str, Any]:
        """Load a Markdown document"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return {
            'content': content,
            'metadata': {
                'source': file_path,
                'file_type': 'markdown',
                'lines': content.count('\n') + 1,
                'characters': len(content)
            }
        }


class HTMLLoader(DocumentLoader):
    """Loader for HTML files"""
    
    name = "HTML Loader"
    description = "Loads HTML files (.html)"
    supported_extensions = ['.html', '.htm']
    required_packages = ['beautifulsoup4']
    
    def load_document(self, file_path: str) -> Dict[str, Any]:
        """Load an HTML document"""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            # Parse HTML and extract text
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.extract()
            
            # Get text
            content = soup.get_text(separator='\n')
            
            return {
                'content': content,
                'metadata': {
                    'source': file_path,
                    'file_type': 'html',
                    'title': soup.title.string if soup.title else None,
                    'characters': len(content)
                }
            }
        except ImportError:
            raise ImportError("beautifulsoup4 is required for HTML loading. Please install it first.")


# Dictionary of available loaders
AVAILABLE_LOADERS = {
    'text': TextLoader,
    'pdf': PDFLoader,
    'docx': DocxLoader,
    'csv': CSVLoader,
    'json': JSONLoader,
    'markdown': MarkdownLoader,
    'html': HTMLLoader
}


def get_available_loaders() -> List[Dict[str, Any]]:
    """Get a list of all available document loaders with their metadata"""
    loaders = []
    
    for loader_id, loader_class in AVAILABLE_LOADERS.items():
        is_available, missing_packages = loader_class.check_dependencies()
        
        loaders.append({
            'id': loader_id,
            'name': loader_class.get_name(),
            'description': loader_class.get_description(),
            'supported_extensions': loader_class.get_supported_extensions(),
            'required_packages': loader_class.get_required_packages(),
            'is_available': is_available,
            'missing_packages': missing_packages
        })
    
    return loaders


def get_loader_for_file(file_path: str) -> Optional[DocumentLoader]:
    """Get the appropriate loader for a file based on its extension"""
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    for loader_class in AVAILABLE_LOADERS.values():
        if ext in loader_class.supported_extensions:
            return loader_class()
    
    return None


def load_document(file_path: str) -> Dict[str, Any]:
    """Load a document using the appropriate loader"""
    loader = get_loader_for_file(file_path)
    
    if loader is None:
        raise ValueError(f"No loader available for file: {file_path}")
    
    # Check dependencies
    is_available, missing_packages = loader.check_dependencies()
    if not is_available:
        raise ImportError(f"Missing required packages: {', '.join(missing_packages)}")
    
    return loader.load_document(file_path)