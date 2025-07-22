from flask import Blueprint, render_template, request, redirect, url_for, current_app, flash, jsonify
from werkzeug.utils import secure_filename
import os
import uuid
import json
import mimetypes
import shutil
# import magic  # Commented out due to Windows compatibility issues
import re
from datetime import datetime
from easy_rag import db
from easy_rag.models import Document, VectorDatabase

# Define allowed file extensions and their MIME types
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'doc', 'csv', 'json', 'html', 'md', 'xml', 'xls', 'xlsx'}

# Define MIME types for validation
ALLOWED_MIME_TYPES = {
    'text/plain': ['.txt', '.md'],
    'application/pdf': ['.pdf'],
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': ['.docx'],
    'application/msword': ['.doc'],
    'text/csv': ['.csv'],
    'application/json': ['.json'],
    'text/html': ['.html'],
    'application/xml': ['.xml'],
    'text/xml': ['.xml'],
    'text/markdown': ['.md'],
    'application/vnd.ms-excel': ['.xls'],
    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': ['.xlsx']
}

def allowed_file(filename):
    """Check if the file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def validate_file(file_path):
    """
    Validate file format and content
    Returns: (is_valid, error_message)
    """
    # Check if file exists
    if not os.path.exists(file_path):
        return False, "File does not exist"
    
    # Check file size (limit to 50MB)
    file_size = os.path.getsize(file_path)
    if file_size > 50 * 1024 * 1024:  # 50MB
        return False, "File size exceeds the maximum limit of 50MB"
    
    # Check if file is empty
    if file_size == 0:
        return False, "File is empty"
    
    # Get file extension
    file_ext = os.path.splitext(file_path)[1].lower()
    
    # Validate MIME type using mimetypes instead of magic
    try:
        mime_type, encoding = mimetypes.guess_type(file_path)
        if not mime_type:
            mime_type = 'application/octet-stream'  # Default MIME type
        
        # Check if MIME type is allowed
        valid_mime = False
        for allowed_mime, extensions in ALLOWED_MIME_TYPES.items():
            if mime_type == allowed_mime and file_ext in extensions:
                valid_mime = True
                break
            # Special case for text files with different encodings
            elif mime_type.startswith('text/') and file_ext in ['.txt', '.md', '.csv', '.json', '.html', '.xml']:
                valid_mime = True
                break
        
        # Special case for files which might have various MIME types
        if file_ext in ['.csv', '.txt', '.pdf', '.docx', '.doc', '.json', '.html', '.md', '.xml', '.xls', '.xlsx']:
            valid_mime = True
        
        if not valid_mime and file_ext not in ALLOWED_EXTENSIONS:
            return False, f"Invalid file format. File extension '{file_ext}' is not allowed."
        
        # Additional validation for specific file types
        if file_ext == '.json':
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    json.load(f)
            except json.JSONDecodeError as e:
                return False, f"Invalid JSON format: {str(e)}"
        
        elif file_ext == '.csv':
            import csv
            # Try different encodings for CSV files
            encodings_to_try = ['utf-8', 'cp949', 'euc-kr', 'latin1']
            success = False
            
            for encoding in encodings_to_try:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        csv_reader = csv.reader(f)
                        # Try to read header row
                        next(csv_reader)
                        success = True
                        break  # Successfully read the file, exit the loop
                except UnicodeDecodeError:
                    continue  # Try the next encoding
                except Exception as e:
                    return False, f"Invalid CSV format: {str(e)}"
            
            if not success:
                return False, f"Invalid CSV format: Could not decode with any of the attempted encodings"
        
        elif file_ext == '.xml':
            try:
                import xml.etree.ElementTree as ET
                ET.parse(file_path)
            except Exception as e:
                return False, f"Invalid XML format: {str(e)}"
        
        elif file_ext == '.xlsx':
            try:
                import openpyxl
                openpyxl.load_workbook(file_path)
            except Exception as e:
                return False, f"Invalid XLSX format: {str(e)}"
        
        elif file_ext == '.xls':
            try:
                import xlrd
                xlrd.open_workbook(file_path)
            except Exception as e:
                return False, f"Invalid XLS format: {str(e)}"
        
        return True, "File is valid"
    
    except Exception as e:
        return False, f"Error validating file: {str(e)}"

bp = Blueprint('document', __name__, url_prefix='/documents')

@bp.route('/')
def index():
    """List all documents"""
    documents = Document.query.all()
    # Convert documents to list of dicts to avoid JSON serialization issues
    documents_dict = [doc.to_dict() for doc in documents]
    return render_template('document/index.html', documents=documents, documents_dict=documents_dict)

def process_uploaded_file(file, extract_metadata=False, description=None):
    """Process an uploaded file and add it to the database"""
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
        
        # Check if file already exists, append timestamp if it does
        if os.path.exists(file_path):
            name, ext = os.path.splitext(filename)
            timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
            filename = f"{name}_{timestamp}{ext}"
            file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
        
        file.save(file_path)
        
        # Validate the file
        is_valid, error_message = validate_file(file_path)
        if not is_valid:
            # Remove the invalid file
            try:
                os.remove(file_path)
            except OSError:
                pass
            
            return False, f'Invalid file: {error_message}'
        
        # Extract metadata if requested
        metadata = {}
        if extract_metadata:
            metadata = extract_document_metadata(file_path)
            
            # Generate preview
            preview_info = generate_document_preview(file_path)
            if preview_info:
                metadata['preview'] = preview_info
        
        # Add description if provided
        if description:
            metadata['description'] = description
        
        # Create document record
        document = Document(
            id=str(uuid.uuid4()),
            path=file_path,
            name=filename,
            type=os.path.splitext(filename)[1].lower(),
            size=os.path.getsize(file_path),
            last_modified=datetime.now(),
            doc_metadata=metadata
        )
        
        db.session.add(document)
        db.session.commit()
        
        return True, document
    else:
        return False, f'File type not allowed. Allowed types: {", ".join(ALLOWED_EXTENSIONS)}'

def process_uploaded_folder(folder_path, extract_metadata=False, description=None):
    """Process all files in an uploaded folder and add them to the database"""
    if not os.path.exists(folder_path):
        return False, f'Folder not found: {folder_path}'
    
    if not os.path.isdir(folder_path):
        return False, f'Not a directory: {folder_path}'
    
    success_count = 0
    error_count = 0
    processed_files = []
    
    # Create a folder in the upload directory
    folder_name = os.path.basename(folder_path)
    timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
    target_folder = os.path.join(current_app.config['UPLOAD_FOLDER'], f"{folder_name}_{timestamp}")
    os.makedirs(target_folder, exist_ok=True)
    
    # Store folder metadata
    folder_metadata = {
        'original_path': folder_path,
        'timestamp': timestamp,
        'description': description
    }
    
    # Process all files in the folder
    for root, _, files in os.walk(folder_path):
        for file in files:
            file_path = os.path.join(root, file)
            rel_path = os.path.relpath(file_path, folder_path)
            target_path = os.path.join(target_folder, rel_path)
            
            # Create subdirectories if needed
            os.makedirs(os.path.dirname(target_path), exist_ok=True)
            
            # Check if file is allowed
            if allowed_file(file):
                try:
                    # Copy file to target directory
                    shutil.copy2(file_path, target_path)
                    
                    # Validate the file
                    is_valid, error_message = validate_file(target_path)
                    if not is_valid:
                        error_count += 1
                        processed_files.append({
                            'name': file,
                            'success': False,
                            'error': error_message
                        })
                        continue
                    
                    # Extract metadata if requested
                    metadata = {}
                    if extract_metadata:
                        metadata = extract_document_metadata(target_path)
                        
                        # Generate preview
                        preview_info = generate_document_preview(target_path)
                        if preview_info:
                            metadata['preview'] = preview_info
                    
                    # Add description if provided
                    if description:
                        metadata['description'] = description
                    
                    # Add folder information to metadata
                    metadata['folder'] = {
                        'name': folder_name,
                        'path': target_folder,
                        'relative_path': rel_path,
                        'parent_dirs': os.path.dirname(rel_path)
                    }
                    
                    # Create document record
                    document = Document(
                        id=str(uuid.uuid4()),
                        path=target_path,
                        name=file,
                        type=os.path.splitext(file)[1].lower(),
                        size=os.path.getsize(target_path),
                        last_modified=datetime.now(),
                        doc_metadata=metadata
                    )
                    
                    db.session.add(document)
                    success_count += 1
                    processed_files.append({
                        'name': file,
                        'success': True,
                        'document_id': document.id
                    })
                except Exception as e:
                    error_count += 1
                    processed_files.append({
                        'name': file,
                        'success': False,
                        'error': str(e)
                    })
            else:
                error_count += 1
                processed_files.append({
                    'name': file,
                    'success': False,
                    'error': f'File type not allowed. Allowed types: {", ".join(ALLOWED_EXTENSIONS)}'
                })
    
    # Commit all changes
    db.session.commit()
    
    return True, {
        'success_count': success_count,
        'error_count': error_count,
        'processed_files': processed_files,
        'folder_path': target_folder,
        'folder_metadata': folder_metadata
    }

@bp.route('/upload', methods=['GET', 'POST'])
def upload():
    """Upload a new document"""
    if request.method == 'POST':
        # Check if the post request has the file part
        if 'file' not in request.files:
            flash('파일이 선택되지 않았습니다.', 'danger')
            return redirect(request.url)
        
        file = request.files['file']
        
        # If user does not select file, browser also
        # submit an empty part without filename
        if file.filename == '':
            flash('파일이 선택되지 않았습니다.', 'danger')
            return redirect(request.url)
        
        # Process file upload
        extract_metadata = request.form.get('extractMetadata') == 'on'
        description = request.form.get('description')
        
        success, result = process_uploaded_file(file, extract_metadata, description)
        if success:
            flash(f'문서 "{result.name}"이 성공적으로 업로드되었습니다.', 'success')
            return redirect(url_for('document.index'))
        else:
            flash(result, 'danger')
            return redirect(request.url)
    
    return render_template('document/upload.html')

@bp.route('/upload-ajax', methods=['POST'])
def upload_ajax():
    """Upload a document via AJAX"""
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'No file part'})
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'success': False, 'error': 'No selected file'})
    
    if file and allowed_file(file.filename):
        try:
            filename = secure_filename(file.filename)
            file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
            
            # Check if file already exists, append timestamp if it does
            if os.path.exists(file_path):
                name, ext = os.path.splitext(filename)
                timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
                filename = f"{name}_{timestamp}{ext}"
                file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
            
            file.save(file_path)
            
            # Validate the file
            is_valid, error_message = validate_file(file_path)
            if not is_valid:
                # Remove the invalid file
                try:
                    os.remove(file_path)
                except OSError:
                    pass
                
                return jsonify({'success': False, 'error': f'Invalid file: {error_message}'})
            
            # Extract metadata if requested
            metadata = {}
            if request.form.get('extractMetadata') == 'true':
                metadata = extract_document_metadata(file_path)
                
                # Generate preview
                preview_info = generate_document_preview(file_path)
                if preview_info:
                    metadata['preview'] = preview_info
            
            # Add description if provided
            if request.form.get('description'):
                metadata['description'] = request.form.get('description')
            
            # Create document record
            document = Document(
                id=str(uuid.uuid4()),
                path=file_path,
                name=filename,
                type=os.path.splitext(filename)[1].lower(),
                size=os.path.getsize(file_path),
                last_modified=datetime.now(),
                doc_metadata=metadata
            )
            
            db.session.add(document)
            db.session.commit()
            
            return jsonify({
                'success': True, 
                'document_id': document.id,
                'document_name': document.name
            })
        except Exception as e:
            return jsonify({'success': False, 'error': str(e)})
    else:
        return jsonify({'success': False, 'error': f'File type not allowed. Allowed types: {", ".join(ALLOWED_EXTENSIONS)}'})

@bp.route('/upload-folder-ajax', methods=['POST'])
def upload_folder_ajax():
    """Upload a folder via AJAX"""
    try:
        if 'files[]' not in request.files:
            return jsonify({'success': False, 'error': '파일이 없습니다.'})
        
        files = request.files.getlist('files[]')
        file_paths = request.form.getlist('file_paths[]')
        folder_name = request.form.get('folder_name')
        
        if not files or not folder_name:
            return jsonify({'success': False, 'error': '폴더 또는 파일이 제공되지 않았습니다.'})
        
        # Extract metadata if requested
        extract_metadata = request.form.get('extractMetadata') == 'true'
        description = request.form.get('description', '')
        
        # Create a folder in the upload directory
        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
        safe_folder_name = secure_filename(folder_name)
        target_folder = os.path.join(current_app.config['UPLOAD_FOLDER'], f"{safe_folder_name}_{timestamp}")
        os.makedirs(target_folder, exist_ok=True)
        
        # Process all files
        success_count = 0
        error_count = 0
        processed_files = []
        
        for i, file in enumerate(files):
            if i >= len(file_paths):
                continue
                
            rel_path = file_paths[i]
            path_parts = rel_path.split('/')
            
            # Create target path
            if len(path_parts) > 1:
                # Create subdirectories if needed
                subdir_parts = path_parts[1:-1]  # Remove folder name and filename
                if subdir_parts:
                    subdir_path = os.path.join(target_folder, *subdir_parts)
                    os.makedirs(subdir_path, exist_ok=True)
                target_path = os.path.join(target_folder, *path_parts[1:])
                relative_path_in_folder = '/'.join(path_parts[1:])
                parent_dirs = '/'.join(path_parts[1:-1]) if len(path_parts) > 2 else ''
            else:
                target_path = os.path.join(target_folder, file.filename)
                relative_path_in_folder = file.filename
                parent_dirs = ''
            
            # Check if file is allowed
            if not allowed_file(file.filename):
                error_count += 1
                processed_files.append({
                    'name': file.filename,
                    'success': False,
                    'error': f'지원되지 않는 파일 형식입니다. 지원 형식: {", ".join(ALLOWED_EXTENSIONS)}'
                })
                continue
            
            try:
                # Save file
                file.save(target_path)
                
                # Validate the file
                is_valid, error_message = validate_file(target_path)
                if not is_valid:
                    # Remove invalid file
                    try:
                        os.remove(target_path)
                    except OSError:
                        pass
                    
                    error_count += 1
                    processed_files.append({
                        'name': file.filename,
                        'success': False,
                        'error': f'파일 검증 실패: {error_message}'
                    })
                    continue
                
                # Extract metadata if requested
                metadata = {}
                if extract_metadata:
                    try:
                        metadata = extract_document_metadata(target_path)
                        
                        # Generate preview
                        preview_info = generate_document_preview(target_path)
                        if preview_info:
                            metadata['preview'] = preview_info
                    except Exception as e:
                        # Continue even if metadata extraction fails
                        current_app.logger.warning(f'메타데이터 추출 실패 for {file.filename}: {str(e)}')
                
                # Add description if provided
                if description:
                    metadata['description'] = description
                
                # Add folder information to metadata
                metadata['folder'] = {
                    'name': folder_name,
                    'path': target_folder,
                    'relative_path': relative_path_in_folder,
                    'parent_dirs': parent_dirs,
                    'upload_timestamp': timestamp
                }
                
                # Create document record
                document = Document(
                    id=str(uuid.uuid4()),
                    path=target_path,
                    name=file.filename,
                    type=os.path.splitext(file.filename)[1].lower(),
                    size=os.path.getsize(target_path),
                    last_modified=datetime.now(),
                    doc_metadata=metadata
                )
                
                db.session.add(document)
                success_count += 1
                processed_files.append({
                    'name': file.filename,
                    'success': True,
                    'document_id': document.id,
                    'path': relative_path_in_folder
                })
                
            except Exception as e:
                error_count += 1
                processed_files.append({
                    'name': file.filename,
                    'success': False,
                    'error': f'파일 처리 중 오류: {str(e)}'
                })
                current_app.logger.error(f'파일 처리 오류 {file.filename}: {str(e)}')
        
        # Commit all changes
        try:
            db.session.commit()
        except Exception as e:
            db.session.rollback()
            return jsonify({
                'success': False,
                'error': f'데이터베이스 저장 중 오류: {str(e)}'
            })
        
        return jsonify({
            'success': True,
            'success_count': success_count,
            'error_count': error_count,
            'processed_files': processed_files,
            'folder_path': target_folder,
            'folder_name': folder_name
        })
        
    except Exception as e:
        current_app.logger.error(f'폴더 업로드 중 예상치 못한 오류: {str(e)}')
        return jsonify({
            'success': False,
            'error': f'업로드 중 오류가 발생했습니다: {str(e)}'
        })

@bp.route('/<id>')
def view(id):
    """View document details"""
    try:
        document = Document.query.get_or_404(id)
        
        # Get vector databases that use this document
        vector_dbs = VectorDatabase.query.filter(
            VectorDatabase.document_ids.contains([document.id])
        ).all()
        
        # Try to get a preview of the document content
        preview_content = None
        preview_error = None
        preview_info = {}  # Initialize as empty dict to avoid None issues
        
        # Check if file exists
        if not os.path.exists(document.path):
            preview_error = f"파일을 찾을 수 없습니다: {document.path}"
        else:
            # Check if preview is already in metadata
            try:
                if document.doc_metadata and isinstance(document.doc_metadata, dict) and 'preview' in document.doc_metadata:
                    preview_info = document.doc_metadata['preview']
                    if isinstance(preview_info, dict):
                        if 'text' in preview_info:
                            preview_content = preview_info['text']
                        if 'error' in preview_info:
                            preview_error = preview_info['error']
                    else:
                        # If preview_info is not a dict, reset it
                        preview_info = {}
            except Exception as e:
                preview_error = f"메타데이터에서 미리보기 정보를 읽는 중 오류 발생: {str(e)}"
                preview_info = {}
            
            # If no preview in metadata or preview_info is empty, generate it on-the-fly
            if not preview_content and not preview_error:
                try:
                    # Check file size before attempting to extract content
                    file_size = os.path.getsize(document.path)
                    if file_size > 10 * 1024 * 1024:  # 10MB
                        preview_error = "파일 크기가 너무 큽니다. 미리보기를 생성할 수 없습니다. (최대 10MB)"
                    else:
                        # Extract content using our utility function
                        content, error = extract_document_content(document.path)
                        
                        if error:
                            preview_error = error
                        elif content:
                            # Truncate content for preview (first 2000 chars)
                            preview_content = content[:2000]
                            if len(content) > 2000:
                                preview_content += "\n\n[내용이 잘렸습니다...]"
                            
                            # Generate full preview info
                            try:
                                preview_info = generate_document_preview(document.path)
                                if not isinstance(preview_info, dict):
                                    preview_info = {}
                                    preview_error = "미리보기 정보 생성 중 오류가 발생했습니다."
                            except Exception as e:
                                # If generate_document_preview fails, create a basic preview_info
                                preview_info = {
                                    'text': preview_content,
                                    'total_length': len(content),
                                    'lines': content.count('\n') + 1,
                                    'words': len(content.split())
                                }
                        else:
                            preview_error = "이 문서에서 내용을 추출할 수 없습니다."
                except Exception as e:
                    preview_error = f"미리보기 로딩 중 오류 발생: {str(e)}"
                    preview_info = {}
        
        # Get file format validation status
        validation_status = None
        try:
            if os.path.exists(document.path):
                is_valid, error_message = validate_file(document.path)
                validation_status = {
                    'is_valid': is_valid,
                    'message': error_message
                }
            else:
                validation_status = {
                    'is_valid': False,
                    'message': f"파일을 찾을 수 없습니다: {document.path}"
                }
        except Exception as e:
            validation_status = {
                'is_valid': False,
                'message': f"파일 검증 중 오류 발생: {str(e)}"
            }
        
        # Convert document to dict to avoid JSON serialization issues
        document_dict = {
            'id': document.id,
            'path': document.path,
            'name': document.name,
            'type': document.type,
            'size': document.size,
            'last_modified': document.last_modified.strftime('%Y-%m-%d %H:%M:%S'),
            'created_at': document.created_at.strftime('%Y-%m-%d %H:%M:%S'),
            'doc_metadata': document.doc_metadata
        }
        
        # Convert vector_dbs to list of dicts to avoid JSON serialization issues
        vector_dbs_dict = []
        for db in vector_dbs:
            vector_dbs_dict.append({
                'id': db.id,
                'name': db.name,
                'vector_store_type': db.vector_store_type,
                'created_at': db.created_at.strftime('%Y-%m-%d %H:%M:%S')
            })
        
        return render_template(
            'document/view.html', 
            document=document,  # Keep original for non-JSON operations
            document_dict=document_dict,  # Use this for JSON serialization
            vector_dbs=vector_dbs,
            vector_dbs_dict=vector_dbs_dict,
            preview_content=preview_content,
            preview_error=preview_error,
            preview_info=preview_info,  # Already ensured to be a dict
            validation_status=validation_status
        )
    except Exception as e:
        # Global exception handler for the entire view function
        return render_template(
            'errors/generic.html',
            error_title="문서 미리보기 오류",
            error_message=f"문서를 표시하는 중 오류가 발생했습니다: {str(e)}",
            back_url=url_for('document.index')
        )

@bp.route('/<id>/delete', methods=['POST'])
def delete(id):
    """Delete a document"""
    try:
        current_app.logger.info(f'Delete request for document ID: {id}')
        document = Document.query.get_or_404(id)
        document_name = document.name
        
        # Check if document is used in any vector databases
        vector_dbs = VectorDatabase.query.filter(
            VectorDatabase.document_ids.contains([document.id])
        ).all()
        
        if vector_dbs:
            # Update vector databases to remove this document ID
            for vdb in vector_dbs:
                vdb.document_ids = [doc_id for doc_id in vdb.document_ids if doc_id != document.id]
                db.session.add(vdb)
        
        # Delete the file
        file_deleted = True
        try:
            if os.path.exists(document.path):
                os.remove(document.path)
        except OSError as e:
            file_deleted = False
            current_app.logger.warning(f'파일 삭제 실패 {document.path}: {str(e)}')
        
        # Delete the database record
        db.session.delete(document)
        db.session.commit()
        
        # Check if this is an AJAX request
        is_ajax = (request.headers.get('X-Requested-With') == 'XMLHttpRequest' or 
                  request.headers.get('Content-Type') == 'application/json' or
                  request.args.get('ajax') == '1')
        
        current_app.logger.info(f'Delete request - AJAX: {is_ajax}, Headers: {dict(request.headers)}')
        
        if is_ajax:
            return jsonify({
                'success': True,
                'message': f'문서 "{document_name}"이 성공적으로 삭제되었습니다.',
                'file_deleted': file_deleted
            })
        else:
            if file_deleted:
                flash(f'문서 "{document_name}"이 성공적으로 삭제되었습니다.', 'success')
            else:
                flash(f'문서 "{document_name}"이 데이터베이스에서 삭제되었지만 파일 삭제에 실패했습니다.', 'warning')
            return redirect(url_for('document.index'))
            
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f'문서 삭제 중 오류: {str(e)}')
        
        if is_ajax:
            return jsonify({
                'success': False,
                'error': f'문서 삭제 중 오류가 발생했습니다: {str(e)}'
            })
        else:
            flash(f'문서 삭제 중 오류가 발생했습니다: {str(e)}', 'danger')
            return redirect(url_for('document.index'))

def extract_document_content(file_path):
    """Extract content from a document file"""
    # Check if file exists
    if not os.path.exists(file_path):
        return None, f"파일을 찾을 수 없습니다: {file_path}"
    
    # Check if file is empty
    if os.path.getsize(file_path) == 0:
        return None, "파일이 비어 있습니다"
    
    # Check if file size is too large (limit to 10MB for preview)
    if os.path.getsize(file_path) > 10 * 1024 * 1024:  # 10MB
        return None, "파일 크기가 너무 큽니다. 미리보기를 생성할 수 없습니다. (최대 10MB)"
    
    file_ext = os.path.splitext(file_path)[1].lower()
    content = None
    error = None
    
    try:
        # Handle different file types
        if file_ext in ['.txt', '.md', '.csv', '.json', '.html', '.xml']:
            # Try different encodings
            encodings_to_try = ['utf-8', 'cp949', 'euc-kr', 'latin1']
            for encoding in encodings_to_try:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break  # Successfully read the file
                except UnicodeDecodeError:
                    continue  # Try the next encoding
                except Exception as e:
                    error = f"파일 읽기 오류: {str(e)}"
                    break
            
            # If content is still None after trying all encodings
            if content is None and error is None:
                error = "지원되는 인코딩으로 파일을 디코딩할 수 없습니다"
            
            # Special handling for very large text content
            if content and len(content) > 1000000:  # If more than ~1MB of text
                content = content[:500000]  # Take first ~500KB
                content += "\n\n[파일이 너무 커서 내용이 잘렸습니다...]"
        
        elif file_ext == '.pdf':
            # First check if the file is actually a PDF
            try:
                with open(file_path, 'rb') as f:
                    header = f.read(5)
                    if header != b'%PDF-':
                        return None, "유효한 PDF 파일이 아닙니다"
            except Exception:
                return None, "PDF 파일 확인 중 오류가 발생했습니다"
                
            try:
                # Try to use PyPDF2 if available
                try:
                    from PyPDF2 import PdfReader
                    reader = PdfReader(file_path)
                    
                    # Limit to first 20 pages for preview
                    max_pages = min(20, len(reader.pages))
                    content = ""
                    
                    for i in range(max_pages):
                        try:
                            page = reader.pages[i]
                            page_text = page.extract_text()
                            if page_text:  # Check if page text extraction was successful
                                content += f"--- 페이지 {i+1} ---\n{page_text}\n\n"
                        except Exception as page_error:
                            # Skip problematic pages
                            content += f"--- 페이지 {i+1} ---\n[내용을 추출할 수 없습니다]\n\n"
                    
                    # If we limited the pages, add a note
                    if len(reader.pages) > max_pages:
                        content += f"\n[미리보기는 처음 {max_pages}페이지만 표시합니다. 전체 문서는 {len(reader.pages)}페이지입니다.]"
                    
                    # If no content was extracted, try another approach or report error
                    if not content.strip():
                        error = "PDF 내용 추출 결과가 비어 있습니다. PDF가 스캔되었거나 이미지만 포함하고 있을 수 있습니다."
                except Exception as pdf_error:
                    error = f"PDF 내용 추출 오류: {str(pdf_error)}"
            except ImportError:
                error = "PyPDF2 라이브러리가 설치되어 있지 않습니다. PDF 내용을 추출할 수 없습니다."
        
        elif file_ext in ['.docx']:
            try:
                # Try to use python-docx if available
                try:
                    import docx
                    doc = docx.Document(file_path)
                    
                    # Extract paragraphs with error handling
                    paragraphs = []
                    for i, paragraph in enumerate(doc.paragraphs):
                        try:
                            if paragraph.text.strip():
                                paragraphs.append(paragraph.text)
                        except Exception:
                            paragraphs.append(f"[단락 {i+1} 읽기 오류]")
                    
                    content = "\n".join(paragraphs)
                    
                    # If no content was extracted, check if document is empty
                    if not content.strip():
                        # Try to extract from tables as well
                        table_content = []
                        for i, table in enumerate(doc.tables):
                            try:
                                for row in table.rows:
                                    row_text = " | ".join([cell.text for cell in row.cells if cell.text.strip()])
                                    if row_text:
                                        table_content.append(row_text)
                            except Exception:
                                table_content.append(f"[테이블 {i+1} 읽기 오류]")
                        
                        if table_content:
                            content = "테이블 내용:\n" + "\n".join(table_content)
                        else:
                            error = "DOCX 내용 추출 결과가 비어 있습니다. 문서가 비어 있거나 이미지/표만 포함하고 있을 수 있습니다."
                except Exception as docx_error:
                    error = f"DOCX 내용 추출 오류: {str(docx_error)}"
            except ImportError:
                error = "python-docx 라이브러리가 설치되어 있지 않습니다. DOCX 내용을 추출할 수 없습니다."
        
        elif file_ext in ['.doc']:
            error = "DOC 형식은 추가 라이브러리가 필요합니다. 현재 지원되지 않습니다."
        
        elif file_ext == '.xlsx':
            try:
                # Check if required libraries are available
                try:
                    import pandas as pd
                    import openpyxl
                    
                    # Use pandas to read Excel file
                    try:
                        # First, get sheet names
                        workbook = openpyxl.load_workbook(file_path, read_only=True)
                        sheet_names = workbook.sheetnames
                        
                        # Read first sheet by default
                        active_sheet = workbook.active.title
                        
                        # Read with pandas (limit rows for preview)
                        df = pd.read_excel(file_path, engine='openpyxl', sheet_name=active_sheet, nrows=100)
                        
                        # Check if DataFrame is empty
                        if df.empty:
                            # Try other sheets
                            for sheet in sheet_names:
                                if sheet != active_sheet:
                                    df = pd.read_excel(file_path, engine='openpyxl', sheet_name=sheet, nrows=100)
                                    if not df.empty:
                                        active_sheet = sheet
                                        break
                            
                            if df.empty:
                                error = "Excel 파일이 비어 있는 것 같습니다"
                                return None, error
                        
                        # Convert DataFrame to string representation (limit columns if too many)
                        if len(df.columns) > 10:
                            content = df.iloc[:, :10].to_string(index=False)
                            content += "\n\n[너무 많은 열이 있어 처음 10개만 표시합니다]"
                        else:
                            content = df.to_string(index=False)
                        
                        # Add sheet information
                        if len(sheet_names) > 1:
                            content = f"시트: {', '.join(sheet_names)}\n\n활성 시트: {active_sheet}\n\n{content}"
                            
                        # Add note if we limited the rows
                        if len(df) == 100:
                            content += "\n\n[미리보기는 처음 100행만 표시합니다]"
                            
                    except Exception as excel_error:
                        error = f"Excel 파일 읽기 오류: {str(excel_error)}"
                except ImportError:
                    error = "pandas 또는 openpyxl이 설치되어 있지 않습니다. XLSX 추출을 위해 먼저 설치하세요."
            except Exception as e:
                error = f"XLSX 내용 추출 오류: {str(e)}"
        
        elif file_ext == '.xls':
            try:
                # Check if required libraries are available
                try:
                    import pandas as pd
                    import xlrd
                    
                    # Use pandas to read Excel file
                    try:
                        # First, get sheet names
                        workbook = xlrd.open_workbook(file_path)
                        sheet_names = workbook.sheet_names()
                        
                        # Read first sheet by default
                        active_sheet = sheet_names[0] if sheet_names else None
                        
                        if not active_sheet:
                            error = "Excel 파일에 시트가 없습니다"
                            return None, error
                        
                        # Read with pandas (limit rows for preview)
                        df = pd.read_excel(file_path, engine='xlrd', sheet_name=active_sheet, nrows=100)
                        
                        # Check if DataFrame is empty
                        if df.empty:
                            # Try other sheets
                            for sheet in sheet_names:
                                if sheet != active_sheet:
                                    df = pd.read_excel(file_path, engine='xlrd', sheet_name=sheet, nrows=100)
                                    if not df.empty:
                                        active_sheet = sheet
                                        break
                            
                            if df.empty:
                                error = "Excel 파일이 비어 있는 것 같습니다"
                                return None, error
                        
                        # Convert DataFrame to string representation (limit columns if too many)
                        if len(df.columns) > 10:
                            content = df.iloc[:, :10].to_string(index=False)
                            content += "\n\n[너무 많은 열이 있어 처음 10개만 표시합니다]"
                        else:
                            content = df.to_string(index=False)
                        
                        # Add sheet information
                        if len(sheet_names) > 1:
                            content = f"시트: {', '.join(sheet_names)}\n\n활성 시트: {active_sheet}\n\n{content}"
                            
                        # Add note if we limited the rows
                        if len(df) == 100:
                            content += "\n\n[미리보기는 처음 100행만 표시합니다]"
                            
                    except Exception as excel_error:
                        error = f"Excel 파일 읽기 오류: {str(excel_error)}"
                except ImportError:
                    error = "pandas 또는 xlrd가 설치되어 있지 않습니다. XLS 추출을 위해 먼저 설치하세요."
            except Exception as e:
                error = f"XLS 내용 추출 오류: {str(e)}"
        
        else:
            error = f"'{file_ext}' 파일 형식은 내용 추출이 지원되지 않습니다."
    
    except Exception as e:
        error = f"내용 추출 오류: {str(e)}"
    
    # Final check to ensure we have either content or error
    if content is None and error is None:
        error = "내용 추출 중 알 수 없는 오류가 발생했습니다"
    
    return content, error

def generate_document_preview(file_path):
    """Generate a preview for a document"""
    preview_info = {}
    
    # Check if file exists
    if not os.path.exists(file_path):
        preview_info['error'] = f"File not found: {file_path}"
        return preview_info
    
    # Check if file is empty
    if os.path.getsize(file_path) == 0:
        preview_info['error'] = "File is empty"
        return preview_info
    
    # Get file extension
    file_ext = os.path.splitext(file_path)[1].lower()
    
    # Extract content
    try:
        content, error = extract_document_content(file_path)
        
        if error:
            preview_info['error'] = error
            return preview_info
        
        if not content:
            preview_info['error'] = "No content could be extracted from this document."
            return preview_info
            
        # Truncate content for preview (first 2000 chars)
        try:
            preview = content[:2000]
            if len(content) > 2000:
                preview += "\n\n[Content truncated...]"
            
            preview_info['text'] = preview
            
            # Add content stats
            preview_info['total_length'] = len(content)
            preview_info['lines'] = content.count('\n') + 1
            preview_info['words'] = len(content.split())
        except Exception as e:
            preview_info['error'] = f"Error processing content: {str(e)}"
            return preview_info
        
        # For structured data, add additional info
        if file_ext == '.json':
            try:
                data = json.loads(content)
                if isinstance(data, dict):
                    preview_info['structure'] = 'object'
                    preview_info['keys'] = list(data.keys())[:10]  # First 10 keys
                    if len(data.keys()) > 10:
                        preview_info['keys'].append('...')
                elif isinstance(data, list):
                    preview_info['structure'] = 'array'
                    preview_info['items'] = len(data)
            except Exception as json_error:
                # Don't fail the entire preview if JSON parsing fails
                preview_info['json_parse_error'] = str(json_error)
        
        elif file_ext == '.csv':
            try:
                import csv
                from io import StringIO
                csv_io = StringIO(content)
                reader = csv.reader(csv_io)
                headers = next(reader, [])
                preview_info['headers'] = headers
                
                # Count rows (up to 100 for performance)
                row_count = 0
                for _ in range(100):
                    try:
                        next(reader)
                        row_count += 1
                    except StopIteration:
                        break
                
                preview_info['sample_rows'] = row_count
                if row_count == 100:
                    preview_info['sample_rows'] = '100+ (sample)'
            except Exception as csv_error:
                # Don't fail the entire preview if CSV parsing fails
                preview_info['csv_parse_error'] = str(csv_error)
                
    except Exception as e:
        preview_info['error'] = f"Error generating preview: {str(e)}"
    
    # Ensure preview_info is a valid dictionary
    if not isinstance(preview_info, dict):
        return {'error': 'Invalid preview information generated'}
    
    return preview_info

def extract_document_metadata(file_path):
    """Extract metadata from a document file"""
    metadata = {}
    
    # Basic file metadata
    file_stat = os.stat(file_path)
    metadata['created'] = datetime.fromtimestamp(file_stat.st_ctime).isoformat()
    metadata['modified'] = datetime.fromtimestamp(file_stat.st_mtime).isoformat()
    
    # Get MIME type
    mime_type, encoding = mimetypes.guess_type(file_path)
    if mime_type:
        metadata['mime_type'] = mime_type
    
    # Extract content-specific metadata based on file type
    file_ext = os.path.splitext(file_path)[1].lower()
    
    try:
        # For text files, count lines, words, characters
        if file_ext in ['.txt', '.md', '.csv', '.json', '.html', '.xml']:
            content = None
            # Try different encodings
            encodings_to_try = ['utf-8', 'cp949', 'euc-kr', 'latin1']
            for encoding in encodings_to_try:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break  # Successfully read the file
                except UnicodeDecodeError:
                    continue  # Try the next encoding
                except Exception as e:
                    metadata['extraction_error'] = f"Error reading file: {str(e)}"
                    break
            
            if content:
                lines = content.count('\n') + 1
                words = len(content.split())
                chars = len(content)
                
                metadata['lines'] = lines
                metadata['words'] = words
                metadata['characters'] = chars
            else:
                metadata['extraction_error'] = "Could not decode file with any of the attempted encodings"
        
        # For JSON files, extract structure info
        if file_ext == '.json':
            json_data = None
            encodings_to_try = ['utf-8', 'cp949', 'euc-kr', 'latin1']
            for encoding in encodings_to_try:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        json_data = json.load(f)
                    if isinstance(json_data, dict):
                        metadata['keys'] = list(json_data.keys())
                    elif isinstance(json_data, list):
                        metadata['items'] = len(json_data)
                    break  # Successfully read the file
                except UnicodeDecodeError:
                    continue  # Try the next encoding
                except json.JSONDecodeError:
                    metadata['json_error'] = 'Invalid JSON format'
                    break
                except Exception as e:
                    metadata['json_error'] = str(e)
                    break
        
        # For CSV files, count rows and columns
        if file_ext == '.csv':
            import csv
            csv_data = None
            encodings_to_try = ['utf-8', 'cp949', 'euc-kr', 'latin1']
            for encoding in encodings_to_try:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        csv_reader = csv.reader(f)
                        headers = next(csv_reader, [])
                        # Store the reader content in a list to count rows
                        rows = list(csv_reader)
                        row_count = len(rows) + 1  # +1 for header
                        
                        metadata['rows'] = row_count
                        metadata['columns'] = len(headers)
                        metadata['headers'] = headers
                        break  # Successfully read the file
                except UnicodeDecodeError:
                    continue  # Try the next encoding
                except Exception as e:
                    metadata['csv_error'] = str(e)
                    break
        
        # For Excel files (XLSX), extract structure info
        if file_ext == '.xlsx':
            try:
                import pandas as pd
                import openpyxl
                
                # Get sheet names
                workbook = openpyxl.load_workbook(file_path, read_only=True)
                sheet_names = workbook.sheetnames
                metadata['sheets'] = sheet_names
                metadata['active_sheet'] = workbook.active.title
                
                # Get data from first sheet
                df = pd.read_excel(file_path, engine='openpyxl')
                metadata['rows'] = len(df)
                metadata['columns'] = len(df.columns)
                metadata['headers'] = df.columns.tolist()
                
            except Exception as e:
                metadata['excel_error'] = str(e)
        
        # For Excel files (XLS), extract structure info
        if file_ext == '.xls':
            try:
                import pandas as pd
                import xlrd
                
                # Get sheet names
                workbook = xlrd.open_workbook(file_path)
                sheet_names = workbook.sheet_names()
                metadata['sheets'] = sheet_names
                metadata['active_sheet'] = sheet_names[0]  # First sheet is default
                
                # Get data from first sheet
                df = pd.read_excel(file_path, engine='xlrd')
                metadata['rows'] = len(df)
                metadata['columns'] = len(df.columns)
                metadata['headers'] = df.columns.tolist()
                
            except Exception as e:
                metadata['excel_error'] = str(e)
    
    except Exception as e:
        metadata['extraction_error'] = str(e)
    
    return metadata